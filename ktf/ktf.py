#! python3
# ------------------------------------------------------------------------------
# the Barbarian Toolsâ„¢
# Kyakuhon Text Formatter
# ------------------------------------------------------------------------------

import enum
import os
import re
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor


# ------------------------------------------------------------------------------

class Decoration(enum.Enum):
    # Decoration
    RESET = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
    REVERSE = '\033[7m'
    # Text Color
    RED = '\033[31m'
    GREEN = '\033[32m'
    BLUE = '\033[34m'
    CYAN = '\033[36m'
    MAGENTA = '\033[35m'
    YELLOW = '\033[33m'
    BLACK = '\033[30m'
    WHITE = '\033[37m'


class Status(enum.Enum):
    MESSAGE = 'ğŸ¤– '
    FAILURE = Decoration.RED.value + 'âŒ [failure] ' + Decoration.RESET.value
    SUCCESS = Decoration.GREEN.value + 'âœ… [success] ' + Decoration.RESET.value
    CAUTION = Decoration.YELLOW.value + 'âš ï¸ [caution] ' + Decoration.RESET.value
    PROSESSING = Decoration.CYAN.value + 'âŒ› [prosessing] ' + Decoration.RESET.value


class LineAttribute(enum.Enum):
    HASHIRA = Decoration.YELLOW.value    # Slugline
    TOGAKI = Decoration.BLUE.value       # Action
    SERIFU = Decoration.GREEN.value      # Dialogue
    BUNRITAI = Decoration.RED.value      # Transition
    MIDASHI = Decoration.CYAN.value
    KAIGYO = ''


def println_col(text: str, col: Decoration):
    print(col.value + text + Decoration.RESET.value)


def message(text: str, status: Status):
    escs = ''
    match status:
        case Status.FAILURE:
            escs = Decoration.RED.value
        case Status.SUCCESS:
            escs = Decoration.GREEN.value
        case Status.CAUTION:
            escs = Decoration.YELLOW.value
        case Status.MESSAGE:
            escs = Decoration.CYAN.value
        case Status.PROSESSING:
            escs = Decoration.CYAN.value
        case _:
            pass

    print(escs + status.value + escs + text + Decoration.RESET.value)


# ------------------------------------------------------------------------------

def load_text_file(file_path: str) -> list[str] | None:
    """ãƒ†ã‚­ã‚¹ãƒˆãƒ•ã‚¡ã‚¤ãƒ«ã‚’èª­ã¿è¾¼ã‚“ã§ã€ãã®æ–‡å­—åˆ—ã‚’è¡Œã”ã¨ã®ãƒªã‚¹ãƒˆã«ã—ã¦è¿”ã™ã€‚"""
    # ãƒ‘ã‚¹ã®å­˜åœ¨ã‚’ç¢ºèª
    if not os.path.isfile(file_path):
        message('There is no file.', Status.FAILURE)
        return None

    # æ‹¡å¼µå­ã‚’ç¢ºèª
    if os.path.splitext(file_path)[1] != '.txt':
        message('This is not a .txt file.', Status.FAILURE)
        return None

    # ãƒ•ã‚¡ã‚¤ãƒ«ã‚’é–‹ã
    with open(file_path, mode='r', encoding='utf_8') as f:
        lines = f.read().splitlines()
        message('Text has been read.', Status.SUCCESS)
        return lines


def text_preprocessor(lines: list[str]) -> list[str]:
    """å„æ–‡å­—åˆ—ã«å‰å‡¦ç†ã‚’åŠ ãˆã¦è¿”ã™ã€‚"""
    # æ–‡å­—åˆ—å‰å¾Œã®ç©ºç™½æ–‡å­—ï¼ˆåŠè§’ãƒ»å…¨è§’ã‚¹ãƒšãƒ¼ã‚¹ï¼ã‚¿ãƒ–ï¼æ”¹è¡Œï¼‰ã‚’å‰Šé™¤
    buffer = [s.strip() for s in lines]
    # ãƒ¡ãƒ¢ï¼šå…ˆé ­ãƒªã‚¹ãƒˆãŒç©ºã®å ´åˆã¨ã‹ã€ä½•ã‹ã‚ã‚Šãã†
    return buffer


def add_attribute_to_line(lines: list[str]) -> list[list[str, LineAttribute]]:
    """è¡Œå˜ä½ã§åˆ¤æ–­ã§ãã‚‹è„šæœ¬å†…å±æ€§ã‚’å„æ–‡å­—åˆ—ã«ä»˜ä¸ã—ãŸå¤šæ¬¡å…ƒãƒªã‚¹ãƒˆã‚’è¿”ã™ã€‚"""
    # æŸ±        ï¼šè¡Œã®å…ˆé ­ã«â—‹â–¡è¨˜å·ãŒå­˜åœ¨
    hashira = re.compile(r'^[â—‹â–¡].+')
    # ã‚»ãƒªãƒ•    ï¼šè¡ŒãŒã€Œã€è¨˜å·ã§å›²ã¾ã‚ŒãŸæ–‡ã§çµ‚äº†
    serifu = re.compile(r'^.*ã€Œ.+ã€$')
    # åˆ†é›¢å¸¯    ï¼šã€ŒÃ—ã€€ã€€ã€€ã€€ã€€Ã—ã€€ã€€ã€€ã€€ã€€Ã—ã€
    bunritai = re.compile(r'^Ã—\s+Ã—\s+Ã—$')
    # è¦‹å‡ºã—    ï¼šè¡ŒãŒã€ã€‘è¨˜å·ã§å›²ã¾ã‚Œã¦ã„ã‚‹
    midashi = re.compile(r'^ã€.+ã€‘$')
    # æ”¹è¡Œï¼šç©ºæ–‡å­—åˆ—ï¼ˆç¾çŠ¶ï¼‰
    # ãƒˆæ›¸ï¼šãã®ä»–ã®æ–‡å­—åˆ—

    line_with_attributes = []
    for l in lines:
        if hashira.match(l) is not None:
            line_with_attributes.append([l, LineAttribute.HASHIRA])
        elif serifu.match(l) is not None:
            line_with_attributes.append([l, LineAttribute.SERIFU])
        elif bunritai.match(l) is not None:
            line_with_attributes.append([('ã€€' + l), LineAttribute.BUNRITAI])
        elif midashi.match(l) is not None:
            line_with_attributes.append([l, LineAttribute.MIDASHI])
        elif l == '':
            line_with_attributes.append([l, LineAttribute.KAIGYO])
        else:
            line_with_attributes.append([l, LineAttribute.TOGAKI])

    # ãƒ¡ãƒ¢ï¼šæœ¬æ–‡ã®å‰ï¼ˆæœ€åˆã®æŸ±ã®å‰ï¼‰ã®æƒ…å ±ã®æ‰±ã„ä¸€è€ƒ

    return line_with_attributes


def fix_line_breaks(line_with_attributes: list[list[str, LineAttribute]]) -> list[list[str, LineAttribute]]:
    """ã‚³ãƒ³ãƒ†ã‚­ã‚¹ãƒˆã‚’åŸºã«æ”¹è¡ŒãŒé©åˆ‡ã‹ãƒã‚§ãƒƒã‚¯ã—ä¿®æ­£ã—ã¦è¿”ã™ã€‚"""
    # ãƒ»æŸ±ã®å‰å¾Œã¯å¸¸ã«ç©ºç™½
    # ãƒ»é•ã†å±æ€§ã¨ã®åˆ‡ã‚Šæ›¿ã‚ã‚Šã«ç©ºç™½

    buffer = []
    context = line_with_attributes[0][1]  # å…ˆé ­ã®æ”¹è¡Œé˜²æ­¢ç­–ã¨ã—ã¦ã¨ã‚Šã‚ãˆãš
    for lwa in line_with_attributes:
        lwa[1]
        if context == lwa[1]:
            # æŸ±ãŒé€£ç¶šã—ã¦ã„ãŸã‚‰æ”¹è¡Œã‚’æŒ¿å…¥ã™ã‚‹ã€‚
            if lwa[1] == LineAttribute.HASHIRA:
                buffer.append(['', LineAttribute.KAIGYO])
        else:
            # æ”¹è¡Œä»¥å¤–ã®é•ã†å±æ€§ãŒé€£ç¶šã—ã¦ã„ãŸã‚‰æ”¹è¡Œã‚’æŒ¿å…¥ã™ã‚‹ã€‚
            if (lwa[1] != LineAttribute.KAIGYO) and (context != LineAttribute.KAIGYO):
                buffer.append(['', LineAttribute.KAIGYO])
        # å†æ§‹ç¯‰
        buffer.append(lwa)
        # ã‚³ãƒ³ãƒ†ã‚¯ã‚¹ãƒˆè¨˜éŒ²
        context = lwa[1]

    return buffer


def format_to_docx(line_with_attributes: list[list[str, LineAttribute]]) -> Document:
    """docxå½¢å¼ã§æ–‡å­—åˆ—ã‚’æ•´å½¢ã—ã¦Documentã‚ªãƒ–ã‚¸ã‚§ã‚¯ãƒˆã‚’è¿”ã™ã€‚"""
    # docxã‚ªãƒ–ã‚¸ã‚§ã‚¯ãƒˆã‚’ä½œæˆ
    doc = Document()
    # ãƒšãƒ¼ã‚¸è¨­å®š
    section = doc.sections[0]
    # ç”¨ç´™ã‚µã‚¤ã‚ºï¼šA4 (210 x 297 mm)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    # ç”¨ç´™ä½™ç™½ï¼š254 mm (1 inch)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    for lwa in line_with_attributes:
        # docxãƒ‘ãƒ©ã‚°ãƒ©ãƒ•ã®è¿½åŠ 
        pgh = None
        if lwa[1] == LineAttribute.HASHIRA:
            pgh = doc.add_heading(lwa[0], 3)
            pgh.style.font.color.rgb = RGBColor(102, 102, 102)
        else:
            pgh = doc.add_paragraph(lwa[0])
        # æ–‡å­—ã®ã‚¹ã‚¿ã‚¤ãƒ«
        pgh.style.font.size = Pt(12)
        pgh.style.font.name = 'Arial'
        # ãƒ‘ãƒ©ã‚°ãƒ©ãƒ•ã®ãƒ•ã‚©ãƒ¼ãƒãƒƒãƒˆ
        pgh_format = pgh.paragraph_format
        pgh_format.space_before = Pt(0)
        pgh_format.space_after = Pt(0)
        pgh_format.line_spacing = Pt(18)  # 1.2pt * 1.5 = 1.8pt

        match lwa[1]:
            case LineAttribute.SERIFU:
                # ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆï¼šï¼‘æ®µ
                pgh_format.left_indent = Inches(0.5)
            case LineAttribute.TOGAKI:
                # ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆï¼šï¼’æ®µ
                pgh_format.left_indent = Inches(1)
            case LineAttribute.BUNRITAI:
                # ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆï¼šï¼’æ®µ
                pgh_format.left_indent = Inches(1)
            case _:
                pass

    return doc


"""
def save_plain_text_file(text: str, path: str):
    f = open(path, encoding='utf_8', mode='w')
    f.write(text)
    f.close()
"""


# ------------------------------------------------------------------------------

def ktf():
    message('Drag and drop the source text file and Enter.', Status.MESSAGE)
    file_path = input()

    lines = load_text_file(file_path)
    # barrier
    if lines == None:
        return None

    lines = text_preprocessor(lines)
    line_with_attributes = add_attribute_to_line(lines)
    line_with_attributes = fix_line_breaks(line_with_attributes)

    """
    for lwa in line_with_attributes:
        println_col(lwa[0], lwa[1])
    """
    message('Generating .docx file. Please wait...', Status.PROSESSING)

    doc = format_to_docx(line_with_attributes)

    dirname = os.path.dirname(file_path)
    basename_without_ext = os.path.splitext(os.path.basename(file_path))[0]
    target_path = os.path.join(dirname, basename_without_ext + '.docx')
    doc.save(target_path)

    message('.docx file Generated at:\n' + target_path, Status.SUCCESS)


def eyecatch():
    """ã‚¢ã‚¤ã‚­ãƒ£ãƒƒãƒ"""
    str = '-----------------------\n'\
          'ğŸ– the Barbarian Toolsâ„¢\n'\
          'Kyakuhon Text Formatter\n'\
          '-----------------------\n'\
          'Beta             v0.1.0\n'\
          '-----------------------'
    println_col(str, Decoration.RED)


# ------------------------------------------------------------------------------

if __name__ == '__main__':
    try:
        eyecatch()
        while True:
            ktf()
            println_col('\n-------Re-Start--------', Decoration.YELLOW)
            message('(Exit: Ctrl + C)', Status.MESSAGE)
    except KeyboardInterrupt:
        println_col('Process KILLED.', Decoration.MAGENTA)
